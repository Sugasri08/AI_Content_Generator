import streamlit as st
from crewai import Agent, Task, Crew
import os

# 📄 PDF + DOCX
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# =========================
# 🎨 PAGE CONFIG
# =========================
st.set_page_config(
    page_title="AI Content Generator",
    page_icon="🤖",
    layout="centered"
)

# =========================
# 🎨 STYLE
# =========================
st.markdown("""
<style>
.stButton>button {
    width: 100%;
    border-radius: 10px;
    height: 45px;
}
</style>
""", unsafe_allow_html=True)

# =========================
# 🔑 CONFIG
# =========================
os.environ["OPENAI_API_KEY"] = os.getenv("GROQ_API_KEY")
os.environ["OPENAI_API_BASE"] = "https://api.groq.com/openai/v1"

# =========================
# 📄 FILE FUNCTIONS
# =========================
def create_pdf(content):
    file_path = "article.pdf"
    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    flowables = []

    for line in content.split("\n"):
        flowables.append(Paragraph(line, styles["Normal"]))

    doc.build(flowables)

    with open(file_path, "rb") as f:
        return f.read()


def create_docx(content):
    file_path = "article.docx"
    doc = Document()
    doc.add_heading("AI Generated Article", 0)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(file_path)

    with open(file_path, "rb") as f:
        return f.read()

# =========================
# 🎯 HEADER
# =========================
st.markdown("<h1 style='text-align: center;'>🤖 AI Content Generator</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center;'>Multi-Agent AI using CrewAI</p>", unsafe_allow_html=True)

st.divider()

# =========================
# 🧠 INPUT
# =========================
topic = st.text_input("💡 Enter a topic:", placeholder="e.g. AI in gaming")

col1, col2 = st.columns(2)
generate = col1.button("🚀 Generate")
clear = col2.button("🧹 Clear")

# =========================
# ⚡ GENERATE
# =========================
if generate and topic:

    with st.spinner("🧠 AI agents are working..."):

        # 🤖 AGENTS
        researcher = Agent(
            role="Researcher",
            goal="Find key points",
            backstory="Fast and accurate",
            llm="groq/llama-3.1-8b-instant"
        )

        writer = Agent(
            role="Writer",
            goal="Write engaging blog",
            backstory="Creative writer",
            llm="groq/llama-3.1-8b-instant"
        )

        editor = Agent(
            role="Editor",
            goal="Polish content",
            backstory="Sharp editor",
            llm="groq/llama-3.1-8b-instant"
        )

        # 📋 TASKS
        research_task = Task(
            description=f"Give 5 short points about {topic}",
            expected_output="Short points",
            agent=researcher
        )

        writing_task = Task(
            description=f"Write a 200-word blog about {topic}",
            expected_output="Short blog",
            agent=writer
        )

        editing_task = Task(
            description="Improve clarity and grammar",
            expected_output="Final blog",
            agent=editor
        )

        crew = Crew(
            agents=[researcher, writer, editor],
            tasks=[research_task, writing_task, editing_task]
        )

        result = crew.kickoff()

    # =========================
    # 📄 OUTPUT
    # =========================
    st.divider()
    st.subheader("✨ Generated Content")

    st.markdown(
        f"""
        <div style="background-color:#0E1117; padding:20px; border-radius:10px;">
        {result.raw}
        </div>
        """,
        unsafe_allow_html=True
    )

    # =========================
    # 📥 DOWNLOAD OPTIONS
    # =========================
    pdf_data = create_pdf(result.raw)
    docx_data = create_docx(result.raw)

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 Download PDF",
            data=pdf_data,
            file_name="article.pdf",
            mime="application/pdf"
        )

    with col2:
        st.download_button(
            label="📝 Download Word",
            data=docx_data,
            file_name="article.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =========================
# 🧹 CLEAR
# =========================
if clear:
    st.rerun()