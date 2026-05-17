from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def create_pdf(content):

    file_path = "generated_content.pdf"

    doc = SimpleDocTemplate(file_path)

    styles = getSampleStyleSheet()

    flowables = []

    for line in content.split("\n"):

        flowables.append(
            Paragraph(line, styles["Normal"])
        )

    doc.build(flowables)

    with open(file_path, "rb") as f:
        return f.read()


def create_docx(content):

    file_path = "generated_content.docx"

    doc = Document()

    doc.add_heading("AI Generated Content", 0)

    for line in content.split("\n"):

        doc.add_paragraph(line)

    doc.save(file_path)

    with open(file_path, "rb") as f:
        return f.read()